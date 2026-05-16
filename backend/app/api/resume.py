from io import BytesIO

from docx import Document
from fastapi import APIRouter, BackgroundTasks, Depends, File, Form, HTTPException, UploadFile
from pypdf import PdfReader
from sqlalchemy.orm import Session

from app.db.session import get_db
from app.models.career import ResumeAnalysis
from app.schemas.career import ResumeAnalyzeRequest, ResumeAnalyzeResponse, ResumeIntelligenceResponse
from app.services.career_service import dumps_list, loads_list
from app.services.realtime_service import realtime_hub
from app.services.resume_service import analyze_resume, build_resume_intelligence

router = APIRouter(prefix="/resume", tags=["resume"])


@router.post("/analyze", response_model=ResumeAnalyzeResponse)
def analyze(payload: ResumeAnalyzeRequest, background_tasks: BackgroundTasks, db: Session = Depends(get_db)):
    result = analyze_resume(payload.resume_text, payload.target_role)
    row = ResumeAnalysis(
        target_role=payload.target_role,
        resume_text=payload.resume_text,
        score=result["score"],
        strengths=dumps_list(result["strengths"]),
        gaps=dumps_list(result["gaps"]),
        suggestions=dumps_list(result["suggestions"]),
    )
    db.add(row)
    db.commit()
    db.refresh(row)
    background_tasks.add_task(
        realtime_hub.broadcast,
        {
            "type": "resume",
            "title": "Resume analysis complete",
            "detail": f"{payload.target_role} resume signal scored {round(row.score)}.",
            "score": row.score,
        },
    )
    return ResumeAnalyzeResponse(
        id=row.id,
        score=row.score,
        strengths=loads_list(row.strengths),
        gaps=loads_list(row.gaps),
        suggestions=loads_list(row.suggestions),
    )


@router.post("/intelligence", response_model=ResumeIntelligenceResponse)
async def intelligence(
    background_tasks: BackgroundTasks,
    file: UploadFile = File(...),
    target_role: str = Form("AI Engineer"),
    fallback_text: str = Form(""),
    db: Session = Depends(get_db),
):
    resume_text = await extract_resume_text(file)
    if len(resume_text.strip()) < 30 and len(fallback_text.strip()) >= 30:
        resume_text = fallback_text
    if len(resume_text.strip()) < 30:
        raise HTTPException(
            status_code=422,
            detail=(
                "Could not extract enough selectable text from this file. "
                "If it is a scanned/image PDF, paste the resume text into the fallback box and run the report again."
            ),
        )

    result = build_resume_intelligence(resume_text, target_role)
    row = ResumeAnalysis(
        target_role=target_role,
        resume_text=resume_text,
        score=result["overall_score"],
        strengths=dumps_list(result["resume_strengths"]),
        gaps=dumps_list([f"Missing or weak signal: {skill}." for skill in result["missing_skills"]]),
        suggestions=dumps_list(result["priority_fixes"]),
    )
    db.add(row)
    db.commit()
    db.refresh(row)
    background_tasks.add_task(
        realtime_hub.broadcast,
        {
            "type": "resume_intelligence",
            "title": "Resume intelligence report ready",
            "detail": f"{file.filename or 'Resume'} matched {result['job_matches'][0]['title']} best.",
            "score": result["overall_score"],
        },
    )
    return ResumeIntelligenceResponse(
        id=row.id,
        file_name=file.filename or "resume",
        target_role=target_role,
        **result,
    )


async def extract_resume_text(file: UploadFile) -> str:
    content = await file.read()
    filename = (file.filename or "").lower()
    content_type = (file.content_type or "").lower()

    if filename.endswith(".txt") or "text/plain" in content_type:
        return content.decode("utf-8", errors="ignore")
    if filename.endswith(".pdf") or "pdf" in content_type:
        return extract_pdf_text(content)
    if filename.endswith(".docx") or "wordprocessingml" in content_type:
        return extract_docx_text(content)
    raise HTTPException(status_code=415, detail="Upload a .txt, .pdf, or .docx resume.")


def extract_pdf_text(content: bytes) -> str:
    try:
        reader = PdfReader(BytesIO(content))
        text_parts = []
        for page in reader.pages:
            text_parts.append(page.extract_text() or "")
            if not text_parts[-1].strip():
                text_parts.append(page.extract_text(extraction_mode="layout") or "")
        return "\n".join(text_parts)
    except Exception as exc:
        raise HTTPException(status_code=422, detail="Could not read text from this PDF.") from exc


def extract_docx_text(content: bytes) -> str:
    try:
        document = Document(BytesIO(content))
        parts = [paragraph.text for paragraph in document.paragraphs]
        for table in document.tables:
            for row in table.rows:
                parts.extend(cell.text for cell in row.cells)
        return "\n".join(parts)
    except Exception as exc:
        raise HTTPException(status_code=422, detail="Could not read text from this DOCX file.") from exc
